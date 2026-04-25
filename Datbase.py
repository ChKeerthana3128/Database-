import streamlit as st
import json
import shutil
import io
from datetime import datetime
from pathlib import Path

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Studio Archive — Interior Design",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── DESIGN TOKENS ────────────────────────────────────────────────────────────
# Deep teal + warm ivory + gold accent palette
BG_COLOR          = "#f0ede6"          # warm ivory
SIDEBAR_BG        = "#1a2e2b"          # deep forest teal
SIDEBAR_TEXT      = "#e8ded0"          # warm ivory text
HERO_GRADIENT     = "linear-gradient(135deg, #0d1f1c 0%, #1a3a34 45%, #2a5a50 100%)"
CARD_BG           = "#ffffff"
ACCENT_GOLD       = "#c9a96e"          # warm gold
ACCENT_TERRA      = "#b87355"          # terracotta
ACCENT_SAGE       = "#6e9e8a"          # sage green
ACCENT_SAND       = "#c4a882"          # warm sand
ACCENT_MOCHA      = "#8c6e56"          # mocha brown
TEXT_DARK         = "#1a1a1a"
TEXT_MID          = "#4a4a4a"
TEXT_LIGHT        = "#7a7a7a"
BORDER_COLOR      = "#ddd6c8"

PHASES = [
    {"key": "phase1", "label": "Phase 1 — Site Visit & CAD Drafting",        "color": ACCENT_TERRA, "icon": "📍"},
    {"key": "phase2", "label": "Phase 2 — Finalising Services & Kitchen",     "color": ACCENT_SAGE,  "icon": "🔧"},
    {"key": "phase3", "label": "Phase 3 — 2D & 3D Designs",                  "color": ACCENT_MOCHA, "icon": "🎨"},
    {"key": "phase4", "label": "Phase 4 — Working Drawings & Selections",     "color": ACCENT_GOLD,  "icon": "📐"},
]

STATUS_META = {
    "Completed":          {"color": ACCENT_SAGE,  "bg": "#eef5f1", "dot": "🟢"},
    "In Progress":        {"color": ACCENT_GOLD,  "bg": "#fdf6e8", "dot": "🟡"},
    "Concept / Proposal": {"color": ACCENT_TERRA, "bg": "#faf0eb", "dot": "🟠"},
    "On Hold":            {"color": ACCENT_MOCHA, "bg": "#f5f0eb", "dot": "⚪"},
}

TYPE_COLORS = [ACCENT_GOLD, ACCENT_SAGE, ACCENT_TERRA, ACCENT_SAND, ACCENT_MOCHA,
               "#b5835a", "#5a9e8a", "#a07850"]

# ─── PATHS ────────────────────────────────────────────────────────────────────
DATA_DIR  = Path("data")
PROJ_FILE = DATA_DIR / "projects.json"
UPLOAD_DIR= DATA_DIR / "uploads"
DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

# ─── GLOBAL CSS ───────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,600;0,700;1,400&family=DM+Sans:wght@300;400;500;600&display=swap');

/* ── Reset & Base ── */
html, body, [class*="css"] {{
    font-family: 'DM Sans', sans-serif;
    color: {TEXT_DARK};
}}

/* ── Main background ── */
.main .block-container {{
    background: {BG_COLOR};
    padding-top: 2rem;
    max-width: 1400px;
}}
.main {{
    background: {BG_COLOR};
}}

/* ── Main content text always dark ── */
.main h1, .main h2, .main h3, .main h4, .main p,
.main label, .main .stMarkdown, .main div {{
    color: #3d2e22 !important;
}}
/* But keep sidebar text light */
section[data-testid="stSidebar"] *,
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] div {{
    color: {SIDEBAR_TEXT} !important;
}}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {{
    background: {SIDEBAR_BG} !important;
    border-right: 1px solid #0d1a18;
}}
section[data-testid="stSidebar"] * {{
    color: {SIDEBAR_TEXT} !important;
}}
section[data-testid="stSidebar"] .stRadio label {{
    color: {SIDEBAR_TEXT} !important;
    font-size: 0.95rem;
    padding: 0.4rem 0;
    cursor: pointer;
}}
section[data-testid="stSidebar"] hr {{
    border-color: #2a4540 !important;
}}
section[data-testid="stSidebar"] .stMarkdown p {{
    color: #a09080 !important;
    font-size: 0.82rem;
}}

/* ── Form Inputs ── */
.stTextInput > div > div > input,
.stNumberInput > div > div > input,
.stTextArea > div > div > textarea {{
    background: #faf8f4 !important;
    border: 1.5px solid {BORDER_COLOR} !important;
    border-radius: 8px !important;
    color: {TEXT_DARK} !important;
    font-family: 'DM Sans', sans-serif !important;
    transition: border-color 0.2s;
}}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {{
    border-color: {ACCENT_GOLD} !important;
    box-shadow: 0 0 0 3px rgba(201,169,110,0.15) !important;
}}
.stSelectbox > div > div,
.stMultiselect > div > div {{
    background: #faf8f4 !important;
    border: 1.5px solid {BORDER_COLOR} !important;
    border-radius: 8px !important;
    color: {TEXT_DARK} !important;
}}
.stSelectbox svg, .stMultiselect svg {{
    color: {ACCENT_GOLD} !important;
}}

/* ── File Uploader ── */
[data-testid="stFileUploadDropzone"],
.stFileUploader > div > div {{
    background: #faf8f4 !important;
    border: 2px dashed {ACCENT_GOLD} !important;
    border-radius: 12px !important;
    color: {TEXT_MID} !important;
    transition: background 0.2s, border-color 0.2s;
}}
[data-testid="stFileUploadDropzone"]:hover {{
    background: #fff8ee !important;
    border-color: {ACCENT_TERRA} !important;
}}

/* ── Submit Button ── */
.stFormSubmitButton > button,
.stButton > button {{
    background: linear-gradient(135deg, {ACCENT_GOLD}, #b8925a) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    font-size: 1rem !important;
    padding: 0.6rem 1.6rem !important;
    transition: transform 0.15s, box-shadow 0.15s !important;
    box-shadow: 0 4px 16px rgba(201,169,110,0.35) !important;
}}
.stFormSubmitButton > button:hover,
.stButton > button:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 22px rgba(201,169,110,0.45) !important;
}}

/* ── Info / Success Alerts ── */
.stAlert {{
    border-radius: 10px !important;
    border-left: 4px solid {ACCENT_GOLD} !important;
}}

/* ── Metric cards ── */
[data-testid="metric-container"] {{
    background: white;
    border-radius: 14px;
    padding: 1rem;
    border: 1px solid {BORDER_COLOR};
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
}}

/* ── Kill ALL Streamlit purple/violet everywhere ── */
:root {{
    --primary-color: {ACCENT_GOLD} !important;
}}

/* Radio buttons — white dots in sidebar */
.stRadio > div > label > div:first-child {{
    border-color: #ffffff !important;
    background: transparent !important;
}}
.stRadio > div > label[data-baseweb="radio"] > div:first-child > div {{
    background: #ffffff !important;
}}
[data-baseweb="radio"] [data-checked="true"] > div {{
    background: #ffffff !important;
    border-color: #ffffff !important;
}}
[data-baseweb="radio"]:focus-within > label > div:first-child {{
    box-shadow: 0 0 0 3px rgba(255,255,255,0.25) !important;
}}

/* File uploader — nuke purple background */
[data-testid="stFileUploader"] {{
    background: transparent !important;
}}
[data-testid="stFileUploadDropzone"] {{
    background: #faf8f2 !important;
    border: 2px dashed #c9a96e !important;
    border-radius: 12px !important;
    color: #4a4a4a !important;
}}
[data-testid="stFileUploadDropzone"]:hover {{
    background: #fff8ee !important;
    border-color: #b87355 !important;
}}
[data-testid="stFileUploadDropzone"] * {{
    color: #4a4a4a !important;
}}
/* The inner wrapper that Streamlit gives a purple tint */
.stFileUploader > div > div,
.stFileUploader > div {{
    background: transparent !important;
    border: none !important;
}}

/* Checkbox */
[data-baseweb="checkbox"] > div {{
    background: {ACCENT_GOLD} !important;
    border-color: {ACCENT_GOLD} !important;
}}

/* Progress / spinner */
.stProgress > div > div {{
    background: {ACCENT_GOLD} !important;
}}

/* Multiselect tags */
[data-baseweb="tag"] {{
    background: #f5f0e8 !important;
    color: #4a4a4a !important;
    border: 1px solid #c9a96e !important;
}}

/* Active page highlight in sidebar radio */
section[data-testid="stSidebar"] [data-baseweb="radio"] [data-checked="true"] > div {{
    background: {ACCENT_GOLD} !important;
}}

/* ── Custom scrollbar ── */
::-webkit-scrollbar {{ width: 7px; }}
::-webkit-scrollbar-track {{ background: {BG_COLOR}; }}
::-webkit-scrollbar-thumb {{ background: {ACCENT_GOLD}; border-radius: 4px; }}

/* ── Helpers ── */
.serif {{ font-family: 'Cormorant Garamond', serif !important; }}
</style>
""", unsafe_allow_html=True)

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def load_projects():
    if PROJ_FILE.exists():
        with open(PROJ_FILE) as f:
            return json.load(f)
    return []

def save_projects(projects):
    with open(PROJ_FILE, "w") as f:
        json.dump(projects, f, indent=2)

def file_icon(name):
    ext = Path(name).suffix.lower()
    m = {
        ".dwg": "📐", ".dxf": "📐",
        ".jpg": "🖼️", ".jpeg": "🖼️", ".png": "🖼️", ".webp": "🖼️",
        ".pdf": "📄", ".mp4": "🎬", ".mov": "🎬",
        ".obj": "🧊", ".skp": "🧊", ".3ds": "🧊",
        ".zip": "🗜️", ".rar": "🗜️",
        ".xlsx": "📊", ".xls": "📊",
        ".docx": "📝", ".doc": "📝",
        ".pptx": "📑",
    }
    return m.get(ext, "📎")

def get_thumbnail(proj_dir: Path):
    if proj_dir.exists():
        for ph in PHASES:
            ph_dir = proj_dir / ph["key"]
            if ph_dir.exists():
                for f in ph_dir.iterdir():
                    if f.suffix.lower() in [".jpg", ".jpeg", ".png", ".webp"]:
                        return f
    return None

def hero(title, subtitle="", icon=""):
    st.markdown(f"""
    <div style="
        background: {HERO_GRADIENT};
        color: white;
        padding: 3rem 3.5rem;
        border-radius: 28px;
        margin-bottom: 2.5rem;
        position: relative;
        overflow: hidden;
    ">
        <div style="
            position: absolute; top: -60px; right: -60px;
            width: 260px; height: 260px;
            background: rgba(201,169,110,0.12);
            border-radius: 50%;
        "></div>
        <div style="
            position: absolute; bottom: -40px; left: 40%;
            width: 180px; height: 180px;
            background: rgba(255,255,255,0.04);
            border-radius: 50%;
        "></div>
        <p style="color:{ACCENT_GOLD}; font-size:0.85rem; letter-spacing:0.18em; text-transform:uppercase; margin:0 0 0.4rem;">
            Studio Archive
        </p>
        <h1 style="
            font-family:'Cormorant Garamond',serif;
            font-size:2.8rem; font-weight:700;
            color:white; margin:0 0 0.5rem; line-height:1.15;
        ">{icon} {title}</h1>
        <p style="color:rgba(255,255,255,0.65); margin:0; font-size:1rem;">{subtitle}</p>
    </div>
    """, unsafe_allow_html=True)

def status_badge(status):
    m = STATUS_META.get(status, {"color": TEXT_MID, "bg": "#eee", "dot": "⚫"})
    return f"""<span style="
        background:{m['bg']}; color:{m['color']};
        padding:3px 12px; border-radius:20px;
        font-size:0.78rem; font-weight:600;
        border:1px solid {m['color']}33;
    ">{m['dot']} {status}</span>"""

def project_card(proj):
    proj_dir = UPLOAD_DIR / proj["id"]
    thumb = get_thumbnail(proj_dir)
    sm = STATUS_META.get(proj.get("status",""), {"color": TEXT_MID, "bg": "#eee", "dot": "⚫"})
    type_color = TYPE_COLORS[hash(proj.get("project_type","")) % len(TYPE_COLORS)]

    # Count files
    total_files = 0
    if proj_dir.exists():
        for p in proj_dir.rglob("*"):
            if p.is_file():
                total_files += 1

    col_img, col_info = st.columns([1, 2.6])
    with col_img:
        if thumb:
            st.image(str(thumb), use_container_width=True)
        else:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, {type_color}22, {type_color}44);
                border-radius: 16px; height: 150px;
                display: flex; align-items: center; justify-content: center;
                font-size: 3rem; border: 1px solid {type_color}33;
            ">🏠</div>
            """, unsafe_allow_html=True)

    with col_info:
        st.markdown(f"""
        <div style="padding: 0.3rem 0;">
            <div style="display:flex; gap:8px; flex-wrap:wrap; margin-bottom:0.6rem;">
                {status_badge(proj.get('status',''))}
                <span style="
                    background:{type_color}18; color:{type_color};
                    padding:3px 12px; border-radius:20px;
                    font-size:0.78rem; font-weight:600;
                    border:1px solid {type_color}33;
                ">{proj.get('project_type','')}</span>
            </div>
            <h3 style="
                font-family:'Cormorant Garamond',serif;
                font-size:1.55rem; font-weight:700;
                margin:0 0 0.2rem; color:{TEXT_DARK};
                line-height:1.2;
            ">{proj.get('name','Untitled')}</h3>
            <p style="color:{TEXT_MID}; margin:0 0 0.5rem; font-size:0.9rem;">
                👤 {proj.get('client','—')} &nbsp;|&nbsp;
                📍 {proj.get('area','—')} &nbsp;|&nbsp;
                🗓️ {proj.get('year','')} &nbsp;|&nbsp;
                📁 {total_files} file{'s' if total_files != 1 else ''}
            </p>
            <p style="color:{TEXT_LIGHT}; font-size:0.88rem; margin:0 0 0.6rem; line-height:1.5;">
                {proj.get('description','')[:140]}{'…' if len(proj.get('description','')) > 140 else ''}
            </p>
        """, unsafe_allow_html=True)

        styles = proj.get("styles", [])
        if styles:
            pills = "".join([
                f'<span style="background:#f0ede6; color:{TEXT_MID}; padding:2px 10px; border-radius:20px; font-size:0.75rem; border:1px solid {BORDER_COLOR}; margin-right:4px;">{s}</span>'
                for s in styles
            ])
            st.markdown(f'<div style="margin-top:0.2rem;">{pills}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

def render_project_detail(proj):
    proj_dir = UPLOAD_DIR / proj["id"]
    sm = STATUS_META.get(proj.get("status",""), {"color": TEXT_MID, "bg": "#eee", "dot": "⚫"})
    type_color = TYPE_COLORS[hash(proj.get("project_type","")) % len(TYPE_COLORS)]

    st.markdown(f"""
    <div style="
        background:white; border-radius:22px;
        padding:2.5rem; margin-bottom:1.5rem;
        box-shadow: 0 8px 36px rgba(0,0,0,0.07);
        border-top: 6px solid {type_color};
    ">
        <div style="display:flex; gap:10px; flex-wrap:wrap; margin-bottom:1rem;">
            {status_badge(proj.get('status',''))}
            <span style="
                background:{type_color}18; color:{type_color};
                padding:3px 14px; border-radius:20px;
                font-size:0.8rem; font-weight:600;
                border:1px solid {type_color}33;
            ">{proj.get('project_type','')}</span>
            <span style="
                background:#f0ede6; color:{TEXT_LIGHT};
                padding:3px 14px; border-radius:20px;
                font-size:0.8rem; border:1px solid {BORDER_COLOR};
            ">{proj.get('budget_range','')}</span>
        </div>
        <h2 style="
            font-family:'Cormorant Garamond',serif;
            font-size:2.2rem; font-weight:700;
            margin:0 0 0.4rem; color:{TEXT_DARK};
        ">{proj.get('name','')}</h2>
        <p style="color:{TEXT_MID}; margin:0 0 1rem;">
            👤 {proj.get('client','—')} &nbsp;·&nbsp;
            📍 {proj.get('area','—')} &nbsp;·&nbsp;
            🗓️ {proj.get('year','')}
        </p>
        <p style="color:{TEXT_MID}; line-height:1.7; margin:0 0 1rem;">
            {proj.get('description','—')}
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Phase files
    st.markdown(f"<h3 style='font-family:Cormorant Garamond,serif; font-size:1.5rem; margin-bottom:1rem;'>📁 Files by Phase</h3>", unsafe_allow_html=True)
    any_files = False
    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
    VIDEO_EXTS = {".mp4", ".mov"}

    for ph in PHASES:
        ph_dir = proj_dir / ph["key"]
        if ph_dir.exists():
            files = sorted(ph_dir.iterdir())
            if files:
                any_files = True

                # Phase header
                st.markdown(f"""
                <div style="
                    background:white; border-radius:16px 16px 0 0;
                    padding:1rem 1.6rem 0.6rem;
                    border-left:5px solid {ph['color']};
                    border-top:1px solid #f0ede6;
                    border-right:1px solid #f0ede6;
                    box-shadow:0 2px 10px rgba(0,0,0,0.04);
                    margin-bottom:0;
                ">
                    <span style="font-weight:600; color:{ph['color']}; font-size:1rem;">
                        {ph['icon']} {ph['label']}
                    </span>
                    <span style="font-weight:400; color:{TEXT_LIGHT}; font-size:0.82rem; margin-left:8px;">
                        {len(files)} file{'s' if len(files)!=1 else ''}
                    </span>
                </div>
                """, unsafe_allow_html=True)

                # Separate images from other files
                images = [f for f in files if f.suffix.lower() in IMAGE_EXTS]
                videos = [f for f in files if f.suffix.lower() in VIDEO_EXTS]
                others = [f for f in files if f.suffix.lower() not in IMAGE_EXTS | VIDEO_EXTS]

                with st.container():
                    st.markdown(f"""
                    <div style="
                        background:white;
                        border-left:5px solid {ph['color']};
                        border-bottom:1px solid #f0ede6;
                        border-right:1px solid #f0ede6;
                        border-radius:0 0 16px 16px;
                        padding:1rem 1.4rem 1.4rem;
                        margin-bottom:1.2rem;
                        box-shadow:0 4px 14px rgba(0,0,0,0.05);
                    ">
                    """, unsafe_allow_html=True)

                    # ── Image Gallery ──
                    if images:
                        st.markdown(f"<p style='font-size:0.82rem; font-weight:600; color:{TEXT_MID}; margin:0.4rem 0 0.6rem;'>🖼️ Images</p>", unsafe_allow_html=True)
                        # Show in rows of 3
                        for i in range(0, len(images), 3):
                            row = images[i:i+3]
                            cols = st.columns(len(row))
                            for col, img_file in zip(cols, row):
                                with col:
                                    st.image(str(img_file), use_container_width=True, caption=img_file.name)
                                    with open(img_file, "rb") as fh:
                                        st.download_button(
                                            "⬇️ Download",
                                            fh.read(),
                                            file_name=img_file.name,
                                            mime=f"image/{img_file.suffix[1:].lower()}",
                                            key=f"dl_{proj['id']}_{ph['key']}_{img_file.name}",
                                            use_container_width=True
                                        )

                    # ── Videos ──
                    if videos:
                        st.markdown(f"<p style='font-size:0.82rem; font-weight:600; color:{TEXT_MID}; margin:0.8rem 0 0.4rem;'>🎬 Videos</p>", unsafe_allow_html=True)
                        for vf in videos:
                            st.markdown(f"<p style='font-size:0.85rem; color:{TEXT_MID}; margin:0.3rem 0;'>{file_icon(vf.name)} {vf.name} <span style='color:{TEXT_LIGHT};'>({round(vf.stat().st_size/1024/1024,1)} MB)</span></p>", unsafe_allow_html=True)
                            st.video(str(vf))
                            with open(vf, "rb") as fh:
                                st.download_button(
                                    "⬇️ Download Video",
                                    fh.read(),
                                    file_name=vf.name,
                                    mime="video/mp4",
                                    key=f"dl_{proj['id']}_{ph['key']}_{vf.name}"
                                )

                    # ── Other Files ──
                    if others:
                        st.markdown(f"<p style='font-size:0.82rem; font-weight:600; color:{TEXT_MID}; margin:0.8rem 0 0.4rem;'>📎 Other Files</p>", unsafe_allow_html=True)
                        for of in others:
                            size_kb = round(of.stat().st_size / 1024, 1)
                            size_str = f"{size_kb} KB" if size_kb < 1024 else f"{round(size_kb/1024,1)} MB"
                            fc1, fc2 = st.columns([3, 1])
                            with fc1:
                                st.markdown(f"""
                                <div style="
                                    display:flex; align-items:center; gap:10px;
                                    padding:8px 12px;
                                    background:#faf8f2; border-radius:8px;
                                    border:1px solid {BORDER_COLOR};
                                    margin-bottom:6px;
                                ">
                                    <span style="font-size:1.4rem;">{file_icon(of.name)}</span>
                                    <div>
                                        <div style="font-size:0.88rem; font-weight:500; color:{TEXT_DARK};">{of.name}</div>
                                        <div style="font-size:0.75rem; color:{TEXT_LIGHT};">{size_str}</div>
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                            with fc2:
                                with open(of, "rb") as fh:
                                    st.download_button(
                                        "⬇️ Download",
                                        fh.read(),
                                        file_name=of.name,
                                        key=f"dl_{proj['id']}_{ph['key']}_{of.name}",
                                        use_container_width=True
                                    )

                    st.markdown("</div>", unsafe_allow_html=True)

    if not any_files:
        st.info("No files uploaded for this project yet.")

    if proj.get("notes"):
        st.markdown(f"""
        <div style="
            background:#fdf8f0; border-radius:14px;
            padding:1.2rem 1.6rem; margin-top:1rem;
            border-left:5px solid {ACCENT_GOLD};
        ">
            <p style="font-weight:600; color:{ACCENT_GOLD}; margin:0 0 0.4rem;">📝 Notes</p>
            <p style="color:{TEXT_MID}; margin:0; line-height:1.7;">{proj.get('notes','')}</p>
        </div>
        """, unsafe_allow_html=True)

def export_excel(proj):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Project Details"
        # Header
        ws["A1"] = "Studio Archive — Project Report"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A2"] = proj.get("name", "")
        ws["A2"].font = Font(bold=True, size=12)
        fields = [
            ("Client", proj.get("client","")),
            ("Year", proj.get("year","")),
            ("Area", proj.get("area","")),
            ("Type", proj.get("project_type","")),
            ("Status", proj.get("status","")),
            ("Budget", proj.get("budget_range","")),
            ("Styles", ", ".join(proj.get("styles",[]))),
            ("Description", proj.get("description","")),
            ("Notes", proj.get("notes","")),
        ]
        for i, (k, v) in enumerate(fields, start=4):
            ws[f"A{i}"] = k
            ws[f"A{i}"].font = Font(bold=True)
            ws[f"B{i}"] = str(v)
        ws.column_dimensions["A"].width = 18
        ws.column_dimensions["B"].width = 50
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()
    except:
        return None

def export_docx(proj):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading(proj.get("name", "Project"), 0)
        doc.add_paragraph(f"Client: {proj.get('client','—')}")
        doc.add_paragraph(f"Year: {proj.get('year','')}")
        doc.add_paragraph(f"Area: {proj.get('area','—')}")
        doc.add_paragraph(f"Type: {proj.get('project_type','')}")
        doc.add_paragraph(f"Status: {proj.get('status','')}")
        doc.add_paragraph(f"Budget: {proj.get('budget_range','')}")
        doc.add_heading("Description", level=2)
        doc.add_paragraph(proj.get("description",""))
        if proj.get("notes"):
            doc.add_heading("Notes", level=2)
            doc.add_paragraph(proj.get("notes",""))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except:
        return None

# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f"""
    <div style="padding: 1.5rem 0 0.5rem;">
        <div style="
            font-family:'Cormorant Garamond',serif;
            font-size:1.9rem; font-weight:700;
            color:{ACCENT_GOLD}; letter-spacing:0.04em;
            line-height:1.1;
        ">🏛️<br>Studio Archive</div>
        <div style="color:#7a9090; font-size:0.78rem; margin-top:0.3rem; letter-spacing:0.1em; text-transform:uppercase;">
            Interior Design Portfolio
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    page = st.radio(
        "Navigation",
        ["📂  All Projects", "➕  Add New Project", "📊  Statistics"],
        label_visibility="collapsed"
    )
    st.markdown("---")

    projects = load_projects()
    completed = sum(1 for p in projects if p.get("status") == "Completed")
    active    = sum(1 for p in projects if p.get("status") == "In Progress")
    concepts  = sum(1 for p in projects if p.get("status") == "Concept / Proposal")
    on_hold   = sum(1 for p in projects if p.get("status") == "On Hold")

    st.markdown(f"""
    <div style="font-size:0.83rem; line-height:2; color:#a09080;">
        <div>📁 <b style="color:{SIDEBAR_TEXT};">{len(projects)}</b> total projects</div>
        <div>✅ <b style="color:{ACCENT_SAGE};">{completed}</b> completed</div>
        <div>🔄 <b style="color:{ACCENT_GOLD};">{active}</b> in progress</div>
        <div>💡 <b style="color:{ACCENT_TERRA};">{concepts}</b> concepts</div>
        <div>⏸️ <b style="color:{ACCENT_SAND};">{on_hold}</b> on hold</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown(f"<div style='color:#506060; font-size:0.75rem; line-height:1.6;'>Studio Archive v2.0<br>© {datetime.now().year} Your Studio</div>", unsafe_allow_html=True)

# ─── PAGE: ALL PROJECTS ───────────────────────────────────────────────────────
if page == "📂  All Projects":
    hero("All Projects", "Your curated interior design portfolio", "📂")

    if not projects:
        st.markdown(f"""
        <div style="
            text-align:center; padding:5rem 2rem;
            background:white; border-radius:22px;
            border:2px dashed {BORDER_COLOR};
        ">
            <div style="font-size:4rem;">🏠</div>
            <h3 style="font-family:'Cormorant Garamond',serif; font-size:1.8rem; color:{TEXT_MID};">No projects yet</h3>
            <p style="color:{TEXT_LIGHT};">Add your first project using the sidebar</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        # ── Filters ──
        col_search, col_type, col_status, col_sort = st.columns([2.5, 1.5, 1.5, 1.5])
        with col_search:
            search = st.text_input("🔍 Search projects…", placeholder="Name, client, area…")
        with col_type:
            types   = ["All Types"] + sorted(set(p.get("project_type","") for p in projects))
            f_type  = st.selectbox("Type", types)
        with col_status:
            statuses = ["All Statuses"] + sorted(set(p.get("status","") for p in projects))
            f_status = st.selectbox("Status", statuses)
        with col_sort:
            sort_by = st.selectbox("Sort by", ["Newest First", "Oldest First", "Name A–Z", "Name Z–A"])

        # Apply filters
        filtered = projects[:]
        if search:
            q = search.lower()
            filtered = [p for p in filtered if
                q in p.get("name","").lower() or
                q in p.get("client","").lower() or
                q in p.get("area","").lower()
            ]
        if f_type != "All Types":
            filtered = [p for p in filtered if p.get("project_type") == f_type]
        if f_status != "All Statuses":
            filtered = [p for p in filtered if p.get("status") == f_status]

        # Sort
        if sort_by == "Newest First":
            filtered = sorted(filtered, key=lambda p: p.get("year",0), reverse=True)
        elif sort_by == "Oldest First":
            filtered = sorted(filtered, key=lambda p: p.get("year",0))
        elif sort_by == "Name A–Z":
            filtered = sorted(filtered, key=lambda p: p.get("name","").lower())
        elif sort_by == "Name Z–A":
            filtered = sorted(filtered, key=lambda p: p.get("name","").lower(), reverse=True)

        st.markdown(f"<p style='color:{TEXT_LIGHT}; font-size:0.88rem; margin-bottom:1.2rem;'>Showing {len(filtered)} of {len(projects)} projects</p>", unsafe_allow_html=True)

        if not filtered:
            st.warning("No projects match your filters.")
        else:
            for proj in filtered:
                type_color = TYPE_COLORS[hash(proj.get("project_type","")) % len(TYPE_COLORS)]
                with st.container():
                    st.markdown(f"""
                    <div style="
                        background:white; border-radius:20px;
                        padding:1.8rem; margin-bottom:1.2rem;
                        box-shadow:0 4px 20px rgba(0,0,0,0.06);
                        border-left:6px solid {type_color};
                        transition: box-shadow 0.2s;
                    ">
                    """, unsafe_allow_html=True)
                    project_card(proj)
                    st.markdown("</div>", unsafe_allow_html=True)

                with st.expander(f"📂 View Full Details — {proj.get('name','')}"):
                    render_project_detail(proj)

                    st.markdown("**Export Project:**")
                    exp_col1, exp_col2, exp_col3 = st.columns(3)
                    with exp_col1:
                        xlsx_data = export_excel(proj)
                        if xlsx_data:
                            st.download_button(
                                "📊 Export Excel",
                                xlsx_data,
                                file_name=f"{proj['id']}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"xlsx_{proj['id']}"
                            )
                    with exp_col2:
                        docx_data = export_docx(proj)
                        if docx_data:
                            st.download_button(
                                "📝 Export Word",
                                docx_data,
                                file_name=f"{proj['id']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"docx_{proj['id']}"
                            )
                    with exp_col3:
                        json_data = json.dumps(proj, indent=2).encode()
                        st.download_button(
                            "📋 Export JSON",
                            json_data,
                            file_name=f"{proj['id']}.json",
                            mime="application/json",
                            key=f"json_{proj['id']}"
                        )

                    st.markdown("---")
                    act_col1, act_col2 = st.columns(2)
                    with act_col1:
                        if st.button(f"✏️ Edit Project", key=f"edit_{proj['id']}", use_container_width=True):
                            st.session_state["editing_id"] = proj["id"]
                            st.rerun()
                    with act_col2:
                        if st.button(f"🗑️ Delete Project", key=f"del_{proj['id']}", use_container_width=True, type="secondary"):
                            projects = [p for p in projects if p["id"] != proj["id"]]
                            save_projects(projects)
                            proj_dir = UPLOAD_DIR / proj["id"]
                            if proj_dir.exists():
                                shutil.rmtree(proj_dir)
                            st.success("Project deleted.")
                            st.rerun()

# ─── PAGE: EDIT PROJECT ───────────────────────────────────────────────────────
if "editing_id" in st.session_state and st.session_state["editing_id"]:
    editing_id = st.session_state["editing_id"]
    projects   = load_projects()
    proj       = next((p for p in projects if p["id"] == editing_id), None)

    if proj:
        hero("Edit Project", f"Updating — {proj.get('name','')}", "✏️")

        if st.button("← Back to All Projects"):
            del st.session_state["editing_id"]
            st.rerun()

        with st.form("edit_form"):
            st.markdown(f"<h3 style='font-family:Cormorant Garamond,serif; color:{TEXT_DARK}; margin-bottom:1rem;'>Project Details</h3>", unsafe_allow_html=True)

            c1, c2 = st.columns(2)
            with c1:
                name         = st.text_input("Project Name *",   value=proj.get("name",""))
                client       = st.text_input("Client Name",       value=proj.get("client",""))
                year         = st.number_input("Year *", min_value=2000, max_value=2035, value=int(proj.get("year", datetime.now().year)))
                area         = st.text_input("Area / Location",   value=proj.get("area",""))
            with c2:
                type_options = ["Residential","Commercial","Office","Hospitality","Retail","Renovation","Other"]
                curr_type    = proj.get("project_type","Residential")
                project_type = st.selectbox("Project Type *", type_options, index=type_options.index(curr_type) if curr_type in type_options else 0)

                budget_options = ["—","Under ₹5L","₹5L–₹20L","₹20L–₹50L","₹50L–₹1Cr","₹1Cr+"]
                curr_budget    = proj.get("budget_range","—")
                budget_range   = st.selectbox("Budget Range", budget_options, index=budget_options.index(curr_budget) if curr_budget in budget_options else 0)

                status_options = ["Completed","In Progress","Concept / Proposal","On Hold"]
                curr_status    = proj.get("status","In Progress")
                status         = st.selectbox("Status", status_options, index=status_options.index(curr_status) if curr_status in status_options else 0)

                all_styles   = ["Contemporary","Modern","Minimalist","Traditional","Luxury","Sustainable","Industrial","Vastu","Japandi","Maximalist"]
                styles       = st.multiselect("Design Styles", all_styles, default=proj.get("styles",[]))

            description = st.text_area("Project Description", value=proj.get("description",""), height=130)

            st.markdown(f"""
            <h3 style='font-family:Cormorant Garamond,serif; color:{TEXT_DARK}; margin:1.5rem 0 0.5rem;'>
                📁 Manage Files by Phase
            </h3>
            <p style='color:{TEXT_LIGHT}; font-size:0.87rem; margin-bottom:0.5rem;'>
                Tick files to delete them, or upload new ones below each phase.
            </p>
            """, unsafe_allow_html=True)

            phase_files    = {}
            files_to_delete = {}

            for ph in PHASES:
                proj_dir = UPLOAD_DIR / proj["id"]
                ph_dir   = proj_dir / ph["key"]
                existing = sorted(ph_dir.iterdir()) if ph_dir.exists() else []

                st.markdown(f"""
                <div style="
                    background: white;
                    border-left: 5px solid {ph['color']};
                    border-radius: 12px;
                    padding: 1rem 1.4rem;
                    margin: 1rem 0 0.4rem;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.04);
                ">
                    <div style="font-weight:600; color:{ph['color']}; margin-bottom:0.6rem;">
                        {ph['icon']} {ph['label']}
                        <span style="font-weight:400; color:{TEXT_LIGHT}; font-size:0.82rem; margin-left:8px;">
                            ({len(existing)} file{'s' if len(existing)!=1 else ''})
                        </span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # Per-file delete checkboxes
                delete_these = []
                if existing:
                    st.markdown(f"<p style='font-size:0.82rem; color:{TEXT_MID}; margin:0 0 0.3rem;'>☑ Tick to delete:</p>", unsafe_allow_html=True)
                    for f in existing:
                        size_kb = round(f.stat().st_size / 1024, 1)
                        checked = st.checkbox(
                            f"{file_icon(f.name)}  {f.name}  ({size_kb} KB)",
                            key=f"del_file_{proj['id']}_{ph['key']}_{f.name}"
                        )
                        if checked:
                            delete_these.append(f)
                else:
                    st.markdown(f"<p style='font-size:0.82rem; color:{TEXT_LIGHT}; font-style:italic;'>No files yet.</p>", unsafe_allow_html=True)

                files_to_delete[ph["key"]] = delete_these

                # Upload new files
                phase_files[ph["key"]] = st.file_uploader(
                    f"➕ Add new files to {ph['label']}",
                    accept_multiple_files=True,
                    key=f"edit_{ph['key']}",
                    label_visibility="collapsed"
                )

            notes = st.text_area("Additional Notes", value=proj.get("notes",""), height=100)

            st.markdown("<br>", unsafe_allow_html=True)
            save_edit = st.form_submit_button("💾  Save Changes", use_container_width=True)

        if save_edit:
            if not name or not name.strip():
                st.error("⚠️ Project Name is required!")
            else:
                # Update fields
                proj["name"]         = name.strip()
                proj["client"]       = client
                proj["year"]         = int(year)
                proj["area"]         = area
                proj["project_type"] = project_type
                proj["budget_range"] = budget_range
                proj["status"]       = status
                proj["styles"]       = styles
                proj["description"]  = description
                proj["notes"]        = notes
                proj["updated_at"]   = datetime.now().isoformat()

                # Delete ticked files
                total_deleted = 0
                for ph in PHASES:
                    for f in files_to_delete.get(ph["key"], []):
                        if f.exists():
                            f.unlink()
                            total_deleted += 1

                # Save new uploaded files
                proj_dir    = UPLOAD_DIR / proj["id"]
                total_new   = 0
                for ph in PHASES:
                    files = phase_files.get(ph["key"]) or []
                    if files:
                        ph_dir = proj_dir / ph["key"]
                        ph_dir.mkdir(parents=True, exist_ok=True)
                        for uf in files:
                            with open(ph_dir / uf.name, "wb") as fh:
                                fh.write(uf.getbuffer())
                            total_new += 1

                # Save back to JSON
                projects = [proj if p["id"] == proj["id"] else p for p in projects]
                save_projects(projects)

                msg = f"✅ **{name}** updated successfully!"
                if total_new:     msg += f"  {total_new} new file{'s' if total_new>1 else ''} added."
                if total_deleted: msg += f"  {total_deleted} file{'s' if total_deleted>1 else ''} deleted."
                st.success(msg)
                del st.session_state["editing_id"]
                st.rerun()

# ─── PAGE: ADD NEW PROJECT ────────────────────────────────────────────────────
elif page == "➕  Add New Project":
    hero("Add New Project", "Document a new interior design commission", "➕")

    with st.form("add_form", clear_on_submit=True):
        st.markdown(f"<h3 style='font-family:Cormorant Garamond,serif; color:{TEXT_DARK}; margin-bottom:1rem;'>Project Details</h3>", unsafe_allow_html=True)

        c1, c2 = st.columns(2)
        with c1:
            name         = st.text_input("Project Name *", placeholder="e.g. Verma Residence — Living Room")
            client       = st.text_input("Client Name", placeholder="e.g. Ravi & Priya Verma")
            year         = st.number_input("Year *", min_value=2000, max_value=2035, value=datetime.now().year)
            area         = st.text_input("Area / Location", placeholder="e.g. Jubilee Hills, Hyderabad")
        with c2:
            project_type = st.selectbox("Project Type *", ["Residential","Commercial","Office","Hospitality","Retail","Renovation","Other"])
            budget_range = st.selectbox("Budget Range", ["—","Under ₹5L","₹5L–₹20L","₹20L–₹50L","₹50L–₹1Cr","₹1Cr+"])
            status       = st.selectbox("Status", ["Completed","In Progress","Concept / Proposal","On Hold"])
            styles       = st.multiselect("Design Styles", ["Contemporary","Modern","Minimalist","Traditional","Luxury","Sustainable","Industrial","Vastu","Japandi","Maximalist"])

        description = st.text_area("Project Description", height=130, placeholder="Describe the design intent, key challenges, and outcomes…")

        st.markdown(f"""
        <h3 style='font-family:Cormorant Garamond,serif; color:{TEXT_DARK}; margin:1.5rem 0 0.5rem;'>
            📁 Upload Files by Phase
        </h3>
        <p style='color:{TEXT_LIGHT}; font-size:0.87rem; margin-bottom:1rem;'>
            Supported: DWG, PDF, JPG, PNG, MP4, OBJ, ZIP, XLSX, DOCX and more.
        </p>
        """, unsafe_allow_html=True)

        phase_files = {}
        for ph in PHASES:
            st.markdown(f"""
            <div style="
                display:flex; align-items:center; gap:8px;
                margin:1rem 0 0.3rem;
                font-weight:600; color:{ph['color']};
            ">
                {ph['icon']} {ph['label']}
            </div>
            """, unsafe_allow_html=True)
            phase_files[ph["key"]] = st.file_uploader(
                f"Upload files for {ph['label']}",
                accept_multiple_files=True,
                key=ph["key"],
                label_visibility="collapsed"
            )

        notes = st.text_area("Additional Notes", height=100, placeholder="Site conditions, client preferences, vendor details…")

        st.markdown("<br>", unsafe_allow_html=True)
        submitted = st.form_submit_button("💾  Save Project", use_container_width=True)

    if submitted:
        if not name or not name.strip():
            st.error("⚠️ Project Name is required!")
        else:
            proj_id = f"{int(datetime.now().timestamp())}_{name[:20].replace(' ','_').lower()}"
            new_proj = {
                "id":           proj_id,
                "name":         name.strip(),
                "client":       client,
                "year":         int(year),
                "project_type": project_type,
                "area":         area,
                "budget_range": budget_range,
                "description":  description,
                "styles":       styles,
                "status":       status,
                "notes":        notes,
                "created_at":   datetime.now().isoformat()
            }

            proj_dir  = UPLOAD_DIR / proj_id
            total_files = 0
            for ph in PHASES:
                files = phase_files.get(ph["key"]) or []
                if files:
                    ph_dir = proj_dir / ph["key"]
                    ph_dir.mkdir(parents=True, exist_ok=True)
                    for uf in files:
                        with open(ph_dir / uf.name, "wb") as fh:
                            fh.write(uf.getbuffer())
                        total_files += 1

            projects = load_projects()
            projects.append(new_proj)
            save_projects(projects)

            st.success(f"✅ **{name}** saved successfully with **{total_files}** file{'s' if total_files != 1 else ''}!")
            st.balloons()

# ─── PAGE: STATISTICS ─────────────────────────────────────────────────────────
elif page == "📊  Statistics":
    hero("Studio Statistics", "An overview of your interior design practice", "📊")

    projects = load_projects()

    if not projects:
        st.info("Add some projects to see your statistics.")
    else:
        # ── Summary Metrics ──
        total    = len(projects)
        done     = sum(1 for p in projects if p.get("status") == "Completed")
        active   = sum(1 for p in projects if p.get("status") == "In Progress")
        concepts = sum(1 for p in projects if p.get("status") == "Concept / Proposal")
        years    = [p.get("year", 0) for p in projects if p.get("year")]

        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Total Projects",   total)
        m2.metric("Completed",        done)
        m3.metric("In Progress",      active)
        m4.metric("Concepts",         concepts)
        m5.metric("Year Range", f"{min(years)}–{max(years)}" if years else "—")

        st.markdown("<br>", unsafe_allow_html=True)

        col_l, col_r = st.columns(2)

        # ── By Type ──
        with col_l:
            st.markdown(f"<h4 style='font-family:Cormorant Garamond,serif; font-size:1.3rem; margin-bottom:1rem;'>Projects by Type</h4>", unsafe_allow_html=True)
            type_counts = {}
            for p in projects:
                t = p.get("project_type","Unknown")
                type_counts[t] = type_counts.get(t,0) + 1

            for i, (t, count) in enumerate(sorted(type_counts.items(), key=lambda x: -x[1])):
                pct  = count / total * 100
                col  = TYPE_COLORS[i % len(TYPE_COLORS)]
                st.markdown(f"""
                <div style="margin-bottom:0.8rem;">
                    <div style="display:flex; justify-content:space-between; margin-bottom:3px;">
                        <span style="font-size:0.9rem; font-weight:500; color:{TEXT_DARK};">{t}</span>
                        <span style="font-size:0.85rem; color:{TEXT_LIGHT};">{count} ({pct:.0f}%)</span>
                    </div>
                    <div style="background:#f0ede6; border-radius:6px; height:10px; overflow:hidden;">
                        <div style="
                            background:linear-gradient(90deg,{col},{col}bb);
                            width:{pct}%; height:100%;
                            border-radius:6px;
                            transition: width 0.6s ease;
                        "></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        # ── By Status ──
        with col_r:
            st.markdown(f"<h4 style='font-family:Cormorant Garamond,serif; font-size:1.3rem; margin-bottom:1rem;'>Projects by Status</h4>", unsafe_allow_html=True)
            status_counts = {}
            for p in projects:
                s = p.get("status","Unknown")
                status_counts[s] = status_counts.get(s,0) + 1

            for s, count in sorted(status_counts.items(), key=lambda x: -x[1]):
                pct = count / total * 100
                sm  = STATUS_META.get(s, {"color": TEXT_MID, "bg": "#eee"})
                st.markdown(f"""
                <div style="margin-bottom:0.8rem;">
                    <div style="display:flex; justify-content:space-between; margin-bottom:3px;">
                        <span style="font-size:0.9rem; font-weight:500; color:{TEXT_DARK};">{s}</span>
                        <span style="font-size:0.85rem; color:{TEXT_LIGHT};">{count} ({pct:.0f}%)</span>
                    </div>
                    <div style="background:#f0ede6; border-radius:6px; height:10px; overflow:hidden;">
                        <div style="
                            background:linear-gradient(90deg,{sm['color']},{sm['color']}bb);
                            width:{pct}%; height:100%;
                            border-radius:6px;
                        "></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        # ── By Year ──
        st.markdown(f"<h4 style='font-family:Cormorant Garamond,serif; font-size:1.3rem; margin:2rem 0 1rem;'>Projects by Year</h4>", unsafe_allow_html=True)
        year_counts = {}
        for p in projects:
            y = str(p.get("year",""))
            if y:
                year_counts[y] = year_counts.get(y, 0) + 1

        if year_counts:
            max_count = max(year_counts.values())
            cols = st.columns(len(year_counts))
            for i, (yr, cnt) in enumerate(sorted(year_counts.items())):
                bar_h = int((cnt / max_count) * 120)
                with cols[i]:
                    st.markdown(f"""
                    <div style="text-align:center;">
                        <div style="
                            height:130px; display:flex;
                            align-items:flex-end; justify-content:center;
                        ">
                            <div style="
                                width:38px; height:{bar_h}px;
                                background:linear-gradient(180deg,{ACCENT_GOLD},{ACCENT_TERRA});
                                border-radius:6px 6px 0 0;
                                position:relative;
                            ">
                                <span style="
                                    position:absolute; top:-22px; left:50%;
                                    transform:translateX(-50%);
                                    font-weight:700; font-size:0.85rem; color:{TEXT_DARK};
                                ">{cnt}</span>
                            </div>
                        </div>
                        <p style="font-size:0.8rem; color:{TEXT_LIGHT}; margin-top:6px;">{yr}</p>
                    </div>
                    """, unsafe_allow_html=True)

        # ── Style cloud ──
        st.markdown(f"<h4 style='font-family:Cormorant Garamond,serif; font-size:1.3rem; margin:2rem 0 1rem;'>Design Style Distribution</h4>", unsafe_allow_html=True)
        all_styles = {}
        for p in projects:
            for s in p.get("styles", []):
                all_styles[s] = all_styles.get(s, 0) + 1

        if all_styles:
            style_html = ""
            max_s = max(all_styles.values())
            for i, (s, cnt) in enumerate(sorted(all_styles.items(), key=lambda x: -x[1])):
                size = 0.8 + (cnt / max_s) * 0.8
                col  = TYPE_COLORS[i % len(TYPE_COLORS)]
                style_html += f'<span style="display:inline-block; margin:4px; padding:5px 14px; background:{col}18; color:{col}; border:1px solid {col}44; border-radius:20px; font-size:{size:.2f}rem; font-weight:500;">{s} <b>({cnt})</b></span>'
            st.markdown(f'<div style="line-height:2.5;">{style_html}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f"<p style='color:{TEXT_LIGHT};'>No styles tagged yet.</p>", unsafe_allow_html=True)
